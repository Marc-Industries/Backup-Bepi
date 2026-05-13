"""Report generation service — LaTeX → PDF via pdflatex, DOCX via python-docx."""
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from dataclasses import dataclass

from jinja2 import Environment, FileSystemLoader

# Lazy import for docxtpl - optional dependency
try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxTemplate = InlineImage = None


TEMPLATE_DIR = Path(__file__).resolve().parents[3] / "templates" / "latex"


def _latex_env() -> Environment:
    return Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        block_start_string=r"\BLOCK{",
        block_end_string="}",
        variable_start_string=r"\VAR{",
        variable_end_string="}",
        comment_start_string=r"\#{",
        comment_end_string="}",
        line_statement_prefix="%%",
        line_comment_prefix="%#",
        trim_blocks=True,
        autoescape=False,
    )


def _escape_latex(text: str) -> str:
    if not text:
        return ""
    chars = {
        "&": r"\&", "%": r"\%", "$": r"\$", "#": r"\#",
        "_": r"\_", "{": r"\{", "}": r"\}",
        "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
    }
    for k, v in chars.items():
        text = text.replace(k, v)
    return text


def _escape_dict(d: dict) -> dict:
    out = {}
    for k, v in d.items():
        if isinstance(v, str):
            out[k] = _escape_latex(v)
        elif isinstance(v, list):
            out[k] = [_escape_dict(x) if isinstance(x, dict) else (_escape_latex(x) if isinstance(x, str) else x) for x in v]
        elif isinstance(v, dict):
            out[k] = _escape_dict(v)
        else:
            out[k] = v
    return out


@dataclass
class ReportResult:
    success: bool
    pdf_path: str | None = None
    log: str = ""
    error: str = ""


def render_latex(template_name: str, context: dict) -> str:
    env = _latex_env()
    tpl = env.get_template(template_name)
    safe_ctx = _escape_dict(context)
    return tpl.render(**safe_ctx)


def compile_pdf(latex_source: str, output_name: str = "report", output_dir: str | None = None) -> ReportResult:
    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="bepi_report_")

    work_dir = tempfile.mkdtemp(prefix="bepi_latex_")
    tex_path = os.path.join(work_dir, f"{output_name}.tex")

    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(latex_source)

    try:
        for _ in range(2):
            result = subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-halt-on-error",
                 f"-output-directory={work_dir}", tex_path],
                capture_output=True, text=True, timeout=60,
                cwd=work_dir,
            )

        pdf_src = os.path.join(work_dir, f"{output_name}.pdf")
        if os.path.exists(pdf_src):
            pdf_dst = os.path.join(output_dir, f"{output_name}.pdf")
            shutil.copy2(pdf_src, pdf_dst)
            return ReportResult(success=True, pdf_path=pdf_dst, log=result.stdout[-2000:] if result.stdout else "")
        else:
            log_path = os.path.join(work_dir, f"{output_name}.log")
            log_content = ""
            if os.path.exists(log_path):
                with open(log_path) as lf:
                    log_content = lf.read()[-3000:]
            return ReportResult(success=False, error="PDF not generated", log=log_content)
    except subprocess.TimeoutExpired:
        return ReportResult(success=False, error="pdflatex timed out (60s)")
    except FileNotFoundError:
        return ReportResult(success=False, error="pdflatex not found — install TeX Live")
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def generate_report(template_name: str, context: dict, output_name: str = "report",
                    output_dir: str | None = None) -> ReportResult:
    latex_source = render_latex(template_name, context)
    return compile_pdf(latex_source, output_name, output_dir)


def generate_docx_report(template_name: str, context: dict, output_name: str = "report",
                         output_dir: str | None = None) -> ReportResult:
    """Generate DOCX report from a template using docxtpl.

    Templates are stored in templates/docx/ as .docx files with Jinja2 placeholders.
    Falls back to creating a document programmatically if template doesn't exist.
    """
    if not DOCX_AVAILABLE:
        return ReportResult(success=False, error="docxtpl not installed. Run: uv sync or pip install docxtpl python-docx")

    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="bepi_report_")

    docx_path = os.path.join(output_dir, f"{output_name}.docx")

    # Try to use template, otherwise build programmatically
    template_path = Path(__file__).resolve().parents[3] / "templates" / "docx" / f"{template_name.replace('.tex', '.docx')}"

    try:
        if template_path.exists():
            doc = DocxTemplate(str(template_path))
            doc.render(context)
            doc.save(docx_path)
        else:
            doc = _build_docx_programmatically(context, template_name)
            doc.save(docx_path)

        return ReportResult(success=True, pdf_path=docx_path)
    except Exception as e:
        return ReportResult(success=False, error=f"DOCX generation failed: {str(e)}")


def _build_docx_programmatically(context: dict, template_name: str):
    """Build a DOCX document programmatically — mirrors the LaTeX/PDF output exactly.

    Structure: Cover → Change Log → TOC → Body (report-specific sections).
    Every Heading 1 is preceded by a page break.
    Numeric columns are right-aligned; totals rows have a distinct background.
    Running header/footer include OOXML PAGE/NUMPAGES fields.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("docxtpl/python-docx not installed")

    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ══════════════════════════════════════════════════════════════════════
    # COLOUR PALETTE  (mirrors ecss_base.tex)
    # ══════════════════════════════════════════════════════════════════════
    ECSS_BLUE    = RGBColor(0x00, 0x33, 0x66)
    ECSS_GRAY    = RGBColor(0x80, 0x80, 0x80)
    ECSS_LIGHT   = RGBColor(0xE6, 0xEC, 0xF5)   # table header fill
    ECSS_TOT     = RGBColor(0xCC, 0xD6, 0xEB)   # totals/subtotals row fill
    ECSS_GREEN   = RGBColor(0x27, 0xAE, 0x60)
    ECSS_RED     = RGBColor(0xE7, 0x4C, 0x3C)
    ECSS_ORANGE  = RGBColor(0xF3, 0x9C, 0x12)

    # ══════════════════════════════════════════════════════════════════════
    # LOW-LEVEL OOXML HELPERS
    # ══════════════════════════════════════════════════════════════════════

    def _set_cell_bg(cell, rgb: RGBColor):
        """Set table-cell background colour via OOXML w:shd."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for existing in tcPr.findall(qn('w:shd')):
            tcPr.remove(existing)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
        tcPr.append(shd)

    def _add_field(run, field_name: str):
        """Insert a Word field code (PAGE / NUMPAGES) into *run*."""
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {field_name} '
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    # ══════════════════════════════════════════════════════════════════════
    # DOCUMENT-LEVEL SETUP
    # ══════════════════════════════════════════════════════════════════════

    def _setup_styles(doc):
        """Configure Calibri as default font; style the heading levels."""
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
        for lvl, sz in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 11)]:
            try:
                st = doc.styles[lvl]
                st.font.name  = 'Calibri'
                st.font.size  = Pt(sz)
                st.font.bold  = True
                st.font.color.rgb = ECSS_BLUE
            except Exception:
                pass

    def _setup_header_footer(doc):
        """
        Running header : [Doc Number] — [Title] — Issue X Rev Y
        Running footer : [Mission] — Page N of M — [Date]
        """
        section = doc.sections[0]

        # ── Header ────────────────────────────────────────────────────────
        hdr = section.header
        hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(
            f"{context.get('doc_number', '')}  —  "
            f"{context.get('doc_title', '')}  —  "
            f"Issue {context.get('issue', '')} Rev {context.get('revision', '')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = ECSS_GRAY

        # ── Footer ────────────────────────────────────────────────────────
        ftr = section.footer
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _fr(text: str | None = None):
            r = fp.add_run(text or "")
            r.font.size = Pt(8)
            r.font.color.rgb = ECSS_GRAY
            return r

        _fr(f"{context.get('mission_name', '')}  —  Page ")
        _add_field(_fr(), 'PAGE')
        _fr(" of ")
        _add_field(_fr(), 'NUMPAGES')
        _fr(f"  —  {context.get('date', '')}")

    # ══════════════════════════════════════════════════════════════════════
    # CONTENT HELPERS
    # ══════════════════════════════════════════════════════════════════════

    def _h1(doc, text: str, *, page_break: bool = True):
        if page_break:
            doc.add_page_break()
        p = doc.add_heading(text, 1)
        for r in p.runs:
            r.font.color.rgb = ECSS_BLUE
        return p

    def _h2(doc, text: str):
        p = doc.add_heading(text, 2)
        for r in p.runs:
            r.font.color.rgb = ECSS_BLUE
        return p

    def _body(doc, text: str):
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].font.size = Pt(11)
        return p

    def _caption(doc, text: str):
        """Bold-italic blue caption placed above a table."""
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = ECSS_BLUE
        return p

    def _make_table(doc, ncols: int, headers: list[str],
                    right_cols: set | None = None):
        """Create a 'Table Grid' table with a styled header row."""
        right_cols = right_cols or set()
        tbl = doc.add_table(rows=1, cols=ncols)
        tbl.style = 'Table Grid'
        hrow = tbl.rows[0]
        for i, hdr in enumerate(headers):
            cell = hrow.cells[i]
            _set_cell_bg(cell, ECSS_LIGHT)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = ECSS_BLUE
            if i in right_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return tbl

    def _add_row(tbl, values: list, *,
                 right_cols: set | None = None,
                 bold: bool = False,
                 bg: RGBColor | None = None):
        """Append a data row; right-align chosen columns; optional bold/bg."""
        right_cols = right_cols or set()
        row = tbl.add_row()
        for i, v in enumerate(values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(v) if v is not None else "—")
            run.font.size = Pt(10)
            run.bold = bold
            if i in right_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if bg:
                _set_cell_bg(cell, bg)
        return row

    def _add_total_row(tbl, values: list, right_cols: set | None = None):
        """Convenience: bold row with ECSS_TOT background (for subtotals/totals)."""
        return _add_row(tbl, values, right_cols=right_cols, bold=True, bg=ECSS_TOT)

    def _status_line(doc, prefix: str, remaining: float):
        """Paragraph: '<prefix>  Status: PASS/FAIL'."""
        p = doc.add_paragraph()
        pr = p.add_run(prefix + "  Status: ")
        pr.font.size = Pt(10)
        sr = p.add_run()
        sr.font.size = Pt(10)
        if remaining > 0:
            sr.text = "PASS"; sr.font.color.rgb = ECSS_GREEN; sr.bold = True
        else:
            sr.text = "FAIL"; sr.font.color.rgb = ECSS_RED;   sr.bold = True

    def _status_run(para, status: str):
        s = (status or "").lower()
        run = para.add_run()
        run.font.size = Pt(10)
        if s == "passed":
            run.text = "PASS"; run.font.color.rgb = ECSS_GREEN; run.bold = True
        elif s == "failed":
            run.text = "FAIL"; run.font.color.rgb = ECSS_RED;   run.bold = True
        elif s == "in_progress":
            run.text = "WIP";  run.font.color.rgb = ECSS_ORANGE; run.bold = True
        else:
            run.text = status or "—"
        return run

    def _bool_run(para, met: bool):
        run = para.add_run()
        run.font.size = Pt(10)
        if met:
            run.text = "PASS"; run.font.color.rgb = ECSS_GREEN; run.bold = True
        else:
            run.text = "FAIL"; run.font.color.rgb = ECSS_RED;   run.bold = True
        return run

    def _risk_run(para, level: str):
        COLOURS = {"critical": ECSS_RED, "high": ECSS_ORANGE,
                   "medium": ECSS_ORANGE, "low": ECSS_GREEN}
        run = para.add_run(level.upper())
        run.font.color.rgb = COLOURS.get(level.lower(), ECSS_GRAY)
        run.bold = True
        run.font.size = Pt(10)
        return run

    def _bullet_list(doc, items: list[str]):
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    def _applicable_docs(doc, items: list[str]):
        _h2(doc, "1.1  Applicable Documents")
        _bullet_list(doc, items)

    # ══════════════════════════════════════════════════════════════════════
    # PAGE 1 — COVER
    # ══════════════════════════════════════════════════════════════════════

    def _build_cover(doc):
        """Cover page — centered text blocks, no tables."""

        def _centered(doc, text, *, size=11, bold=False, color=None, italic=False, space_after=None):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if space_after is not None:
                p.paragraph_format.space_after = Pt(space_after)
            r = p.add_run(text)
            r.font.size  = Pt(size)
            r.font.bold  = bold
            r.font.name  = 'Calibri'
            if color:
                r.font.color.rgb = color
            if italic:
                r.italic = True
            return p

        # ── Decorative rule ───────────────────────────────────────────
        _centered(doc, "─" * 90, size=11, bold=True, color=ECSS_BLUE)

        # ── Title & subtitle ──────────────────────────────────────────
        _centered(doc, context.get('doc_title', 'Report'), size=26, bold=True, color=ECSS_BLUE)
        _centered(doc, context.get('doc_subtitle', ''), size=14, color=ECSS_GRAY)

        # ── Decorative rule ───────────────────────────────────────────
        _centered(doc, "─" * 90, size=11, bold=True, color=ECSS_BLUE)

        # ── Metadata — centred label : value pairs ────────────────────
        meta_items = [
            ("Document Number", context.get('doc_number', '')),
            ("Issue",           context.get('issue', '1')),
            ("Revision",        context.get('revision', '0')),
            ("Date",            context.get('date', '')),
            ("Status",          context.get('doc_status', 'Draft')),
            ("Mission",         context.get('mission_name', '')),
            ("Phase",           context.get('phase', '')),
            ("Customer",        context.get('customer', '')),
            ("Prime Contractor", context.get('prime_contractor', '')),
        ]
        for label, value in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            rl = p.add_run(f"{label}: ")
            rl.font.size = Pt(11)
            rl.font.bold = True
            rl.font.color.rgb = ECSS_BLUE
            rv = p.add_run(str(value))
            rv.font.size = Pt(11)

        # ── Signatures Table ──────────────────────────────────────────
        doc.add_paragraph()  # spacer
        sig_tbl = doc.add_table(rows=4, cols=4)
        sig_tbl.style = 'Table Grid'
        
        # Header
        for ci, hdr in enumerate(["", "Name", "Date", "Signature"]):
            cell = sig_tbl.rows[0].cells[ci]
            _set_cell_bg(cell, ECSS_LIGHT)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = ECSS_BLUE
            
        # Rows
        for ri, (role, person, date_val) in enumerate([
            ("Prepared by", context.get('prepared_by', ''), context.get('date', '')),
            ("Reviewed by", context.get('reviewed_by', ''), ''),
            ("Approved by", context.get('approved_by', ''), ''),
        ], start=1):
            row = sig_tbl.rows[ri]
            for ci, val in enumerate([role, person, date_val, ""]):
                p = row.cells[ci].paragraphs[0]
                p.clear()
                run = p.add_run(str(val))
                run.font.size = Pt(10)
                if ci == 0:
                    run.bold = True

        # ── Footer note ───────────────────────────────────────────────
        _centered(doc, "Generated by B.E.P.I. — Budget, Engineering & Project Integration",
                  size=9, color=ECSS_GRAY, italic=True)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # PAGE 2 — DOCUMENT CHANGE LOG
    # ══════════════════════════════════════════════════════════════════════

    def _build_change_log(doc):
        h = doc.add_heading("Document Change Log", 1)
        for r in h.runs:
            r.font.color.rgb = ECSS_BLUE

        cl = doc.add_table(rows=2, cols=4)
        cl.style = 'Table Grid'
        for ci, hdr in enumerate(["Issue", "Rev", "Date", "Description"]):
            cell = cl.rows[0].cells[ci]
            _set_cell_bg(cell, ECSS_LIGHT)
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = ECSS_BLUE
        row = cl.rows[1]
        for ci, val in enumerate([
            context.get('issue', '1'),
            context.get('revision', '0'),
            context.get('date', ''),
            "Initial issue",
        ]):
            row.cells[ci].paragraphs[0].add_run(val).font.size = Pt(10)

        # Removed page break to keep TOC on the same page


    # ══════════════════════════════════════════════════════════════════════
    # PAGE 3 — TABLE OF CONTENTS  (OOXML TOC field + manual fallback)
    # ══════════════════════════════════════════════════════════════════════

    def _build_toc(doc, entries: list[tuple[str, str]]):
        """Insert a real Word TOC field code so Word / LibreOffice can update
        page numbers automatically.  A manual fallback listing is added so
        the document is still readable before the field is refreshed."""
        h = doc.add_heading("Table of Contents", 2)
        for r in h.runs:
            r.font.color.rgb = ECSS_BLUE

        # ── Real OOXML TOC field ──────────────────────────────────────
        p_toc = doc.add_paragraph()
        run_toc = p_toc.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run_toc._r.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
        run_toc._r.append(instr)

        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')
        run_toc._r.append(fld_separate)

        # Fallback text (visible until user presses "Update field" in Word)
        for num, title in entries:
            is_sub = '.' in num
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20) if is_sub else Pt(0)
            r_num = p.add_run(f"{num}  ")
            r_num.font.size = Pt(10 if is_sub else 11)
            r_num.bold = not is_sub
            r_num.font.color.rgb = ECSS_GRAY if is_sub else ECSS_BLUE
            r_title = p.add_run(title)
            r_title.font.size = Pt(10 if is_sub else 11)
            r_title.bold = not is_sub
            r_title.font.color.rgb = ECSS_GRAY if is_sub else ECSS_BLUE

        # Close the TOC field
        p_end = doc.add_paragraph()
        run_end = p_end.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run_end._r.append(fld_end)

        # Removed page break; Section 1 will provide its own page break via _h1


    # ══════════════════════════════════════════════════════════════════════
    # REPORT BODY — BUDGET REPORT  (mirrors budget_report.tex exactly)
    # ══════════════════════════════════════════════════════════════════════

    def _build_budget_report(doc):
        sys_margin  = context.get('system_margin', 0)
        RC_MASS_SUB = {2, 3, 4, 5}   # Nominal, w/Margin, Margin%, Items
        RC_PWR_SUB  = {2, 3, 4}      # Nominal, w/Margin, Margin%
        RC_EQ       = {2, 3, 4, 5}   # Nom, Qty, Margin%, Total (equipment)

        # ── 1 Scope ───────────────────────────────────────────────────────
        _h1(doc, "1  Scope")
        _body(doc,
            f"This document presents the mass and power budget for the "
            f"{context.get('mission_name', 'N/A')} mission at Phase "
            f"{context.get('phase', 'N/A')}. Budget allocations are computed "
            f"in accordance with ECSS-E-HB-10-02A margin philosophy."
        )
        _applicable_docs(doc, [
            "ECSS-E-HB-10-02A — Verification guidelines",
            "ECSS-M-ST-10C — Project planning and implementation",
            f"{context.get('mission_name', '')} System Requirements Document "
            f"({context.get('doc_number', '').replace('BUD', 'SRD')})",
        ])

        # ── 2 Margin Philosophy ───────────────────────────────────────────
        _h1(doc, "2  Margin Philosophy")
        _body(doc,
            "Margins are applied according to ECSS-E-HB-10-02A, Table 1, "
            "as a function of mission phase and equipment maturity."
        )
        margin_table = context.get('margin_table', [])
        _caption(doc, f"Table 1 — Component Margin Policy — Phase {context.get('phase', '')}")
        mt1 = _make_table(doc, 4,
                          ["Phase", "Estimate (%)", "Measured (%)", "Qualified (%)"],
                          right_cols={1, 2, 3})
        for row in margin_table:
            _add_row(mt1, [
                row.get('phase', ''),
                f"{row.get('estimate', '')}%",
                f"{row.get('measured', '')}%",
                f"{row.get('qualified', '')}%",
            ], right_cols={1, 2, 3})
        _body(doc,
            f"System-level margin for Phase {context.get('phase', '')}: {sys_margin}%."
        )

        # ── 3 Mass Budget ─────────────────────────────────────────────────
        _h1(doc, "3  Mass Budget")
        _h2(doc, "3.1  Summary")
        _caption(doc, "Table 2 — Mass Budget Summary")
        mt2 = _make_table(doc, 6,
                          ["Subsystem", "Code",
                           "Nominal [kg]", "w/ Margin [kg]", "Margin [%]", "Items"],
                          right_cols=RC_MASS_SUB)
        for s in context.get('mass_by_subsystem', []):
            _add_row(mt2, [
                s.get('name', ''), s.get('code', ''),
                f"{s.get('nominal', 0):.1f}",
                f"{s.get('with_margin', 0):.1f}",
                f"{s.get('margin_pct', 0):.1f}",
                str(s.get('count', 0)),
            ], right_cols=RC_MASS_SUB)

        # Totals rows
        _add_total_row(mt2,
            ["Dry Mass Subtotal", "",
             f"{context.get('dry_nominal', 0):.1f}",
             f"{context.get('dry_with_margin', 0):.1f}", "", ""],
            right_cols=RC_MASS_SUB)
        _add_row(mt2,
            [f"System Margin ({sys_margin}%)", "",
             "", f"{context.get('system_margin_kg', 0):.1f}", "", ""],
            right_cols=RC_MASS_SUB, bg=ECSS_TOT)
        _add_total_row(mt2,
            ["Dry Mass with System Margin", "",
             "", f"{context.get('dry_with_system', 0):.1f}", "", ""],
            right_cols=RC_MASS_SUB)
        _add_row(mt2,
            ["Propellant", "",
             f"{context.get('propellant_kg', 0):.1f}",
             f"{context.get('propellant_kg', 0):.1f}", "", ""],
            right_cols=RC_MASS_SUB, bg=ECSS_TOT)
        _add_total_row(mt2,
            ["Wet Mass", "",
             "", f"{context.get('wet_mass', 0):.1f}", "", ""],
            right_cols=RC_MASS_SUB)

        _status_line(doc,
            f"Budget Limit: {context.get('mass_limit', 0):.1f} kg  |  "
            f"Remaining: {context.get('mass_remaining', 0):.1f} kg",
            context.get('mass_remaining', 0))

        # 3.2 Equipment level breakdown
        _h2(doc, "3.2  Mass Breakdown — Equipment Level")
        _caption(doc, "Table 3 — Mass Equipment Breakdown")
        mt3 = _make_table(doc, 7,
                          ["Code", "Name", "Nom. [kg]", "Qty",
                           "Margin [%]", "Total [kg]", "Mat."],
                          right_cols=RC_EQ)
        for e in context.get('mass_equipment', []):
            _add_row(mt3, [
                e.get('code', ''), e.get('name', ''),
                f"{e.get('nominal', 0):.2f}", str(e.get('qty', 1)),
                f"{e.get('margin_pct', 0):.0f}", f"{e.get('total', 0):.2f}",
                e.get('maturity', ''),
            ], right_cols=RC_EQ)

        # ── 4 Power Budget ────────────────────────────────────────────────
        _h1(doc, "4  Power Budget")
        _h2(doc, "4.1  Summary — Nominal Mode")
        _caption(doc, "Table 4 — Power Budget Summary — Nominal Mode")
        pt4 = _make_table(doc, 5,
                          ["Subsystem", "Code",
                           "Nominal [W]", "w/ Margin [W]", "Margin [%]"],
                          right_cols=RC_PWR_SUB)
        for s in context.get('power_by_subsystem', []):
            _add_row(pt4, [
                s.get('name', ''), s.get('code', ''),
                f"{s.get('nominal', 0):.1f}",
                f"{s.get('with_margin', 0):.1f}",
                f"{s.get('margin_pct', 0):.1f}",
            ], right_cols=RC_PWR_SUB)

        _add_total_row(pt4,
            ["Total", "",
             f"{context.get('power_nominal', 0):.1f}",
             f"{context.get('power_with_margin', 0):.1f}", ""],
            right_cols=RC_PWR_SUB)
        _add_row(pt4,
            [f"System Margin ({sys_margin}%)", "",
             "", f"{context.get('power_system_margin_w', 0):.1f}", ""],
            right_cols=RC_PWR_SUB, bg=ECSS_TOT)
        _add_total_row(pt4,
            ["Total with System Margin", "",
             "", f"{context.get('power_with_system', 0):.1f}", ""],
            right_cols=RC_PWR_SUB)

        _status_line(doc,
            f"Budget Limit: {context.get('power_limit', 0):.1f} W  |  "
            f"Remaining: {context.get('power_remaining', 0):.1f} W",
            context.get('power_remaining', 0))

        # 4.2 Equipment level breakdown
        _h2(doc, "4.2  Power Breakdown — Equipment Level")
        _caption(doc, "Table 5 — Power Equipment Breakdown")
        pt5 = _make_table(doc, 7,
                          ["Code", "Name", "Nom. [W]", "Qty",
                           "Margin [%]", "Total [W]", "Mat."],
                          right_cols=RC_EQ)
        for e in context.get('power_equipment', []):
            _add_row(pt5, [
                e.get('code', ''), e.get('name', ''),
                f"{e.get('nominal', 0):.1f}", str(e.get('qty', 1)),
                f"{e.get('margin_pct', 0):.0f}", f"{e.get('total', 0):.1f}",
                e.get('maturity', ''),
            ], right_cols=RC_EQ)

    # ══════════════════════════════════════════════════════════════════════
    # REPORT BODY — REQUIREMENTS DOCUMENT  (mirrors requirements_doc.tex)
    # ══════════════════════════════════════════════════════════════════════

    def _build_requirements_doc(doc):

        def _req_block(req: dict):
            _h2(doc, f"{req.get('req_id', '')} — {req.get('title', '')}")
            rows_data = [
                ("Text:",         req.get('text', '')),
                ("Category:",     req.get('category', '').capitalize()),
                ("Parent:",       req.get('parent_id', '') or "N/A"),
                ("V&V Method:",   (req.get('method', '') or 'TBD').upper()),
                ("Status:",       req.get('status', '')),
                ("Allocated to:", req.get('allocated_to', '') or "N/A"),
                ("Owner:",        req.get('owner', '') or "TBD"),
            ]
            rt = doc.add_table(rows=len(rows_data), cols=2)
            rt.style = 'Table Grid'
            for i, (label, value) in enumerate(rows_data):
                row = rt.rows[i]
                _set_cell_bg(row.cells[0], ECSS_LIGHT)
                lr = row.cells[0].paragraphs[0].add_run(label)
                lr.bold = True
                lr.font.size = Pt(10)
                vp = row.cells[1].paragraphs[0]
                if label == "Status:":
                    _status_run(vp, value)
                else:
                    vr = vp.add_run(str(value))
                    vr.font.size = Pt(10)

        _h1(doc, "1  Scope")
        _body(doc,
            f"This document defines the requirements for the "
            f"{context.get('mission_name', 'N/A')} mission at Phase "
            f"{context.get('phase', 'N/A')}, in accordance with ECSS-E-ST-10C."
        )
        _applicable_docs(doc, [
            "ECSS-E-ST-10C — System engineering general requirements",
            "ECSS-E-ST-10-06C — Technical requirements specification",
            f"{context.get('mission_name', '')} Mission Definition Document",
        ])

        _h1(doc, "2  Requirements Overview")
        cov = context.get('coverage_by_level', [])
        _caption(doc, "Table 1 — Requirements Coverage Summary")
        ct = _make_table(doc, 5,
                         ["Level", "Total", "Verified", "In Progress", "Coverage (%)"],
                         right_cols={1, 2, 3, 4})
        for lvl in cov:
            _add_row(ct, [
                str(lvl.get('level', '')).capitalize(),
                str(lvl.get('total', 0)),
                str(lvl.get('verified', 0)),
                str(lvl.get('in_progress', 0)),
                f"{lvl.get('pct', 0):.0f}%",
            ], right_cols={1, 2, 3, 4})
        _add_total_row(ct, [
            "Total",
            str(context.get('total_reqs', 0)),
            str(context.get('total_verified', 0)),
            str(context.get('total_in_progress', 0)),
            f"{context.get('overall_pct', 0):.0f}%",
        ], right_cols={1, 2, 3, 4})

        for sec_n, (title, key) in enumerate([
            ("Stakeholder Requirements", "reqs_stakeholder"),
            ("Mission Requirements",     "reqs_mission"),
            ("System Requirements",      "reqs_system"),
            ("Subsystem Requirements",   "reqs_subsystem"),
            ("Equipment Requirements",   "reqs_equipment"),
        ], start=3):
            reqs = context.get(key, [])
            if not reqs:
                continue
            _h1(doc, f"{sec_n}  {title}")
            for req in reqs:
                _req_block(req)

        _h1(doc, "Verification Matrix")
        _caption(doc, "Verification Matrix — all requirements")
        vm = _make_table(doc, 6,
                         ["ID", "Title", "Level", "Method", "Status", "Owner"])
        for req in context.get('all_reqs', []):
            row = vm.add_row()
            for ci, v in enumerate([
                req.get('req_id', ''),
                req.get('title', ''),
                req.get('level', ''),
                (req.get('method', '') or 'TBD').upper(),
                None,
                req.get('owner', ''),
            ]):
                p = row.cells[ci].paragraphs[0]
                if ci == 4:
                    _status_run(p, req.get('status', ''))
                else:
                    p.add_run(str(v) if v else "—").font.size = Pt(10)

        _h1(doc, "Traceability Matrix")
        _caption(doc, "Traceability Matrix")
        tr = _make_table(doc, 3, ["Requirement", "Parent", "Allocated To"])
        for req in context.get('all_reqs', []):
            _add_row(tr, [
                req.get('req_id', ''),
                req.get('parent_id', '') or "—",
                req.get('allocated_to', '') or "—",
            ])

    # ══════════════════════════════════════════════════════════════════════
    # REPORT BODY — RISK ASSESSMENT  (mirrors risk_assessment.tex)
    # ══════════════════════════════════════════════════════════════════════

    def _build_risk_assessment(doc):
        _h1(doc, "1  Scope")
        _body(doc,
            f"This document presents the risk assessment for "
            f"{context.get('mission_name', 'N/A')} at Phase "
            f"{context.get('phase', 'N/A')}, in accordance with ECSS-M-ST-80C."
        )
        _applicable_docs(doc, [
            "ECSS-M-ST-80C — Risk Management",
            "ECSS-Q-ST-30C — Dependability",
            "ECSS-Q-ST-40C — Safety",
        ])

        _h1(doc, "2  Risk Assessment Methodology")
        _h2(doc, "2.1  Risk Matrix")
        _body(doc, "Risks are assessed on a 5×5 matrix (Likelihood × Consequence):")
        _caption(doc, "Likelihood Scale")
        lt = _make_table(doc, 2, ["Score", "Likelihood"], right_cols={0})
        for score, desc in [
            ("1", "Rare — unlikely to occur"),
            ("2", "Unlikely — could occur but not expected"),
            ("3", "Possible — might occur"),
            ("4", "Likely — will probably occur"),
            ("5", "Almost certain — expected to occur"),
        ]:
            _add_row(lt, [score, desc], right_cols={0})

        _caption(doc, "Consequence Scale")
        ct = _make_table(doc, 2, ["Score", "Consequence"], right_cols={0})
        for score, desc in [
            ("1", "Negligible"), ("2", "Minor"), ("3", "Moderate"),
            ("4", "Major"),      ("5", "Catastrophic"),
        ]:
            _add_row(ct, [score, desc], right_cols={0})

        _body(doc, "Risk Level: Critical (≥15), High (≥9), Medium (≥4), Low (<4).")

        _h2(doc, "2.2  Risk Summary")
        _caption(doc, "Risk Summary by Level")
        rs = _make_table(doc, 4, ["Level", "Count", "Open", "Mitigating"],
                         right_cols={1, 2, 3})
        for lvl in context.get('risk_summary', []):
            _add_row(rs, [
                str(lvl.get('level', '')).capitalize(),
                str(lvl.get('count', 0)),
                str(lvl.get('open', 0)),
                str(lvl.get('mitigating', 0)),
            ], right_cols={1, 2, 3})
        _add_total_row(rs, [
            "Total",
            str(context.get('total_risks', 0)),
            str(context.get('total_open', 0)),
            str(context.get('total_mitigating', 0)),
        ], right_cols={1, 2, 3})

        _h1(doc, "3  Risk Register")
        for risk in context.get('risks', []):
            _h2(doc, f"{risk.get('risk_id', '')} — {risk.get('title', '')}")
            rows_data = [
                ("Description:", risk.get('description', '')),
                ("Category:",    str(risk.get('category', '')).capitalize()),
                ("Likelihood:",  str(risk.get('likelihood', ''))),
                ("Consequence:", str(risk.get('consequence', ''))),
                ("Risk Score:",  None),
                ("Status:",      risk.get('status', '')),
                ("Owner:",       risk.get('owner', '')),
                ("Mitigation:",  risk.get('mitigation', '')),
            ]
            if risk.get('residual_l'):
                rows_data.append(("Residual L/C:",
                    f"{risk['residual_l']}/{risk.get('residual_c', '')} "
                    f"(Score: {risk['residual_l'] * risk.get('residual_c', 0)})"
                ))
            rk = doc.add_table(rows=len(rows_data), cols=2)
            rk.style = 'Table Grid'
            for i, (label, value) in enumerate(rows_data):
                row = rk.rows[i]
                _set_cell_bg(row.cells[0], ECSS_LIGHT)
                lr = row.cells[0].paragraphs[0].add_run(label)
                lr.bold = True
                lr.font.size = Pt(10)
                vp = row.cells[1].paragraphs[0]
                if label == "Risk Score:":
                    vp.add_run(f"{risk.get('score', 0)} (")
                    _risk_run(vp, risk.get('level', ''))
                    vp.add_run(")")
                elif label == "Status:":
                    _status_run(vp, value)
                else:
                    vp.add_run(str(value) if value else "—").font.size = Pt(10)

        _h1(doc, "4  FMECA Summary")
        _caption(doc, "FMECA Summary — Top 10 by RPN")
        fm = _make_table(doc, 7,
                         ["Node", "Failure Mode", "S", "O", "D", "RPN", "Mitigation"],
                         right_cols={2, 3, 4, 5})
        for f in context.get('fmeca', []):
            row = fm.add_row()
            for ci, v in enumerate([
                f.get('node', ''), f.get('mode', ''),
                str(f.get('severity', '')), str(f.get('occurrence', '')),
                str(f.get('detection', '')), str(f.get('rpn', '')),
                f.get('mitigation', ''),
            ]):
                p = row.cells[ci].paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(10)
                if ci == 5:
                    run.bold = True
                if ci in {2, 3, 4, 5}:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ══════════════════════════════════════════════════════════════════════
    # REPORT BODY — DESIGN DEFINITION FILE  (mirrors design_definition.tex)
    # ══════════════════════════════════════════════════════════════════════

    def _build_design_definition(doc):
        _h1(doc, "1  Scope")
        _body(doc,
            f"This Design Definition File (DDF) describes the system design for "
            f"{context.get('mission_name', 'N/A')} at Phase "
            f"{context.get('phase', 'N/A')}, in accordance with ECSS-E-ST-10C."
        )
        _applicable_docs(doc, [
            "ECSS-E-ST-10C — System engineering general requirements",
            "ECSS-E-HB-10-02A — Verification guidelines",
            f"{context.get('mission_name', '')} System Requirements Document",
        ])

        _h1(doc, "2  Mission Overview")
        _caption(doc, "Mission Overview")
        mo = doc.add_table(rows=5, cols=2)
        mo.style = 'Table Grid'
        for i, (label, key) in enumerate([
            ("Mission",       "mission_name"),
            ("Orbit",         "orbit"),
            ("Phase",         "phase"),
            ("Target Launch", "launch_date"),
            ("Lifetime",      "lifetime"),
        ]):
            row = mo.rows[i]
            _set_cell_bg(row.cells[0], ECSS_LIGHT)
            lr = row.cells[0].paragraphs[0].add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            val = context.get(key, "")
            if key == "phase":
                val = f"{val} — {context.get('phase_name', '')}"
            row.cells[1].paragraphs[0].add_run(str(val)).font.size = Pt(10)

        _h1(doc, "3  Product Tree")
        _h2(doc, "3.1  System Architecture")
        _caption(doc, "Product Tree")
        pt = _make_table(doc, 5, ["Code", "Name", "Level", "Parent", "Subsystem"])
        for n in context.get('product_tree', []):
            _add_row(pt, [
                n.get('code', ''), n.get('name', ''), n.get('level', ''),
                n.get('parent', '') or "—", n.get('subsystem', ''),
            ])

        _h1(doc, "4  Subsystem Descriptions")
        for sub in context.get('subsystems', []):
            _h2(doc, f"{sub.get('name', '')} ({sub.get('code', '')})")
            sd = doc.add_table(rows=5, cols=2)
            sd.style = 'Table Grid'
            for i, (label, value) in enumerate([
                ("Components:",      sub.get('components', '')),
                ("Mass (nominal):",  f"{sub.get('mass', 0):.1f} kg"),
                ("Power (nominal):", f"{sub.get('power', 0):.1f} W"),
                ("ECSS Standard:",   sub.get('ecss_standard', '')),
                ("Lead:",            sub.get('lead', 'TBD')),
            ]):
                row = sd.rows[i]
                _set_cell_bg(row.cells[0], ECSS_LIGHT)
                lr = row.cells[0].paragraphs[0].add_run(label)
                lr.bold = True
                lr.font.size = Pt(10)
                row.cells[1].paragraphs[0].add_run(str(value)).font.size = Pt(10)

        _h1(doc, "5  Budget Summary")
        _h2(doc, "5.1  Mass Budget")
        _caption(doc, "Mass Budget Summary")
        mb = doc.add_table(rows=5, cols=3)
        mb.style = 'Table Grid'
        for i, (label, nom, mar) in enumerate([
            ("Dry Mass",      f"{context.get('dry_nominal', 0):.1f}",   f"{context.get('dry_with_margin', 0):.1f}"),
            ("System Margin", "",                                         f"{context.get('system_margin_kg', 0):.1f}"),
            ("Dry Mass Total","",                                         f"{context.get('dry_with_system', 0):.1f}"),
            ("Propellant",    f"{context.get('propellant_kg', 0):.1f}", f"{context.get('propellant_kg', 0):.1f}"),
            ("Wet Mass",      "",                                         f"{context.get('wet_mass', 0):.1f}"),
        ]):
            row = mb.rows[i]
            row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
            for ci, v in [(1, nom), (2, mar)]:
                p = row.cells[ci].paragraphs[0]
                p.add_run(v).font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _status_line(doc,
            f"Limit: {context.get('mass_limit', 0):.1f} kg",
            context.get('mass_remaining', 0))

        _h2(doc, "5.2  Power Budget")
        _caption(doc, "Power Budget Summary")
        pb = doc.add_table(rows=3, cols=3)
        pb.style = 'Table Grid'
        for i, (label, nom, mar) in enumerate([
            ("Equipment Total", f"{context.get('power_nominal', 0):.1f}",
             f"{context.get('power_with_margin', 0):.1f}"),
            ("System Margin",   "",
             f"{context.get('power_system_margin_w', 0):.1f}"),
            ("Total",           "",
             f"{context.get('power_with_system', 0):.1f}"),
        ]):
            row = pb.rows[i]
            row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
            for ci, v in [(1, nom), (2, mar)]:
                p = row.cells[ci].paragraphs[0]
                p.add_run(v).font.size = Pt(10)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _status_line(doc,
            f"Limit: {context.get('power_limit', 0):.1f} W",
            context.get('power_remaining', 0))

        _h1(doc, "6  Key Design Drivers")
        _bullet_list(doc, context.get('design_drivers', []))

        _h1(doc, "7  Open Items & TBDs")
        _caption(doc, "Open Items")
        oi = _make_table(doc, 3, ["#", "Type", "Description"])
        for i, item in enumerate(context.get('open_items', []), start=1):
            _add_row(oi, [str(i), item.get('type', ''), item.get('description', '')])

    # ══════════════════════════════════════════════════════════════════════
    # REPORT BODY — REVIEW DATA PACKAGE  (mirrors review_data_package.tex)
    # ══════════════════════════════════════════════════════════════════════

    def _build_review_data_package(doc):
        _h1(doc, "1  Scope")
        _body(doc,
            f"This Review Data Package supports the "
            f"{context.get('review_name', 'N/A')} for "
            f"{context.get('mission_name', 'N/A')}, Phase "
            f"{context.get('phase', 'N/A')}."
        )
        _h2(doc, "1.1  Review Information")
        ri = doc.add_table(rows=4, cols=2)
        ri.style = 'Table Grid'
        for i, (label, value) in enumerate([
            ("Review",           context.get('review_name', '')),
            ("Phase Transition", f"{context.get('phase_before', '')} → {context.get('phase_after', '')}"),
            ("Planned Date",     context.get('review_date', '')),
            ("Status",           context.get('review_status', '')),
        ]):
            row = ri.rows[i]
            _set_cell_bg(row.cells[0], ECSS_LIGHT)
            lr = row.cells[0].paragraphs[0].add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(str(value)).font.size = Pt(10)

        _h1(doc, "2  Entry Criteria Assessment")
        _caption(doc, "Entry Criteria")
        ec = _make_table(doc, 3, ["#", "Criterion", "Status"])
        for i, c in enumerate(context.get('entry_criteria', []), start=1):
            row = ec.add_row()
            row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(c.get('text', '')).font.size = Pt(10)
            _bool_run(row.cells[2].paragraphs[0], c.get('met', False))

        _h1(doc, "3  Deliverables Status")
        _caption(doc, "Deliverables")
        dl = _make_table(doc, 5, ["DRD Code", "Title", "Status", "Owner", "Due"])
        for d in context.get('deliverables', []):
            _add_row(dl, [
                d.get('drd_code', ''), d.get('title', ''),
                str(d.get('status', '')).upper(),
                d.get('owner', ''), d.get('due_date', ''),
            ])

        _h1(doc, "4  Technical Status Summary")
        _h2(doc, "4.1  Requirements Coverage")
        rc = _make_table(doc, 4, ["Level", "Total", "Verified", "Coverage (%)"],
                         right_cols={1, 2, 3})
        for lvl in context.get('req_coverage', []):
            _add_row(rc, [
                str(lvl.get('level', '')).capitalize(),
                str(lvl.get('total', 0)),
                str(lvl.get('verified', 0)),
                f"{lvl.get('pct', 0):.0f}%",
            ], right_cols={1, 2, 3})

        _h2(doc, "4.2  Budget Status")
        _caption(doc, "Budget Status")
        bs = _make_table(doc, 5, ["Budget", "Allocated", "Limit", "Margin", "Status"])
        for label, alloc, limit, margin, remaining in [
            ("Mass (wet)",
             f"{context.get('wet_mass', 0):.1f} kg",
             f"{context.get('mass_limit', 0):.1f} kg",
             f"{context.get('mass_remaining', 0):.1f} kg",
             context.get('mass_remaining', 0)),
            ("Power",
             f"{context.get('power_with_system', 0):.1f} W",
             f"{context.get('power_limit', 0):.1f} W",
             f"{context.get('power_remaining', 0):.1f} W",
             context.get('power_remaining', 0)),
        ]:
            row = bs.add_row()
            for ci, v in enumerate([label, alloc, limit, margin]):
                row.cells[ci].paragraphs[0].add_run(v).font.size = Pt(10)
            _bool_run(row.cells[4].paragraphs[0], remaining > 0)

        _h2(doc, "4.3  Risk Status")
        _caption(doc, "Risk Counts by Level")
        rsk = _make_table(doc, 5, ["", "Critical", "High", "Medium", "Low"],
                          right_cols={1, 2, 3, 4})
        _add_row(rsk, [
            "Count",
            str(context.get('risk_critical', 0)),
            str(context.get('risk_high', 0)),
            str(context.get('risk_medium', 0)),
            str(context.get('risk_low', 0)),
        ], right_cols={1, 2, 3, 4})

        _h2(doc, "4.4  Schedule Status")
        _caption(doc, "Schedule Status")
        sched = _make_table(doc, 2, ["Metric", "Value"])
        for label, value in [
            ("Project Duration", f"{context.get('schedule_duration', 0)} days"),
            ("End Date",         str(context.get('schedule_end', ''))),
            ("Critical Tasks",   str(context.get('critical_tasks', 0))),
            ("Overall Progress", f"{context.get('overall_progress', 0):.0f}%"),
        ]:
            _add_row(sched, [label, value])

        _h1(doc, "5  Top Risks")
        for risk in context.get('top_risks', []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{risk.get('risk_id', '')}: ").bold = True
            p.add_run(
                f"{risk.get('title', '')} "
                f"(L={risk.get('likelihood', '')}, C={risk.get('consequence', '')}, "
                f"Score={risk.get('score', '')}) — "
                f"{str(risk.get('status', '')).upper()}"
            ).font.size = Pt(11)

        _h1(doc, "6  Action Items")
        _caption(doc, "Action Items")
        ai = _make_table(doc, 6, ["#", "Type", "Action", "Owner", "Due", "Status"])
        for i, a in enumerate(context.get('action_items', []), start=1):
            _add_row(ai, [
                str(i), a.get('type', ''), a.get('action', ''),
                a.get('owner', ''), a.get('due', ''), a.get('status', ''),
            ])

        _h1(doc, "7  Recommendation")
        _body(doc, context.get('recommendation', ''))

    # ══════════════════════════════════════════════════════════════════════
    # TOC DEFINITIONS (one per report type)
    # ══════════════════════════════════════════════════════════════════════

    TOC_BUDGET = [
        ("1",   "Scope"),
        ("1.1", "Applicable Documents"),
        ("2",   "Margin Philosophy"),
        ("3",   "Mass Budget"),
        ("3.1", "Summary"),
        ("3.2", "Mass Breakdown — Equipment Level"),
        ("4",   "Power Budget"),
        ("4.1", "Summary — Nominal Mode"),
        ("4.2", "Power Breakdown — Equipment Level"),
    ]
    TOC_REQUIREMENTS = [
        ("1",   "Scope"),
        ("1.1", "Applicable Documents"),
        ("2",   "Requirements Overview"),
        ("3",   "Stakeholder Requirements"),
        ("4",   "Mission Requirements"),
        ("5",   "System Requirements"),
        ("6",   "Subsystem Requirements"),
        ("7",   "Equipment Requirements"),
        ("8",   "Verification Matrix"),
        ("9",   "Traceability Matrix"),
    ]
    TOC_RISK = [
        ("1",   "Scope"),
        ("1.1", "Applicable Documents"),
        ("2",   "Risk Assessment Methodology"),
        ("2.1", "Risk Matrix"),
        ("2.2", "Risk Summary"),
        ("3",   "Risk Register"),
        ("4",   "FMECA Summary"),
    ]
    TOC_DDF = [
        ("1",   "Scope"),
        ("1.1", "Applicable Documents"),
        ("2",   "Mission Overview"),
        ("3",   "Product Tree"),
        ("3.1", "System Architecture"),
        ("4",   "Subsystem Descriptions"),
        ("5",   "Budget Summary"),
        ("5.1", "Mass Budget"),
        ("5.2", "Power Budget"),
        ("6",   "Key Design Drivers"),
        ("7",   "Open Items & TBDs"),
    ]
    TOC_RDP = [
        ("1",   "Scope"),
        ("1.1", "Review Information"),
        ("2",   "Entry Criteria Assessment"),
        ("3",   "Deliverables Status"),
        ("4",   "Technical Status Summary"),
        ("4.1", "Requirements Coverage"),
        ("4.2", "Budget Status"),
        ("4.3", "Risk Status"),
        ("4.4", "Schedule Status"),
        ("5",   "Top Risks"),
        ("6",   "Action Items"),
        ("7",   "Recommendation"),
    ]

    # ══════════════════════════════════════════════════════════════════════
    # ASSEMBLE THE DOCUMENT
    # ══════════════════════════════════════════════════════════════════════

    doc = Document()
    _setup_styles(doc)
    _setup_header_footer(doc)

    # Page 1 — Cover
    _build_cover(doc)

    # Page 2 — Change Log
    _build_change_log(doc)

    # Page 3 — TOC + Body
    tpl = template_name.lower()
    if 'budget' in tpl:
        _build_toc(doc, TOC_BUDGET)
        _build_budget_report(doc)
    elif 'requirement' in tpl:
        _build_toc(doc, TOC_REQUIREMENTS)
        _build_requirements_doc(doc)
    elif 'risk' in tpl:
        _build_toc(doc, TOC_RISK)
        _build_risk_assessment(doc)
    elif 'design' in tpl or 'ddf' in tpl:
        _build_toc(doc, TOC_DDF)
        _build_design_definition(doc)
    elif 'review' in tpl or 'rdp' in tpl:
        _build_toc(doc, TOC_RDP)
        _build_review_data_package(doc)
    else:
        _h1(doc, context.get('doc_title', 'Report'), page_break=False)
        _body(doc, "No specific template builder matched. Check template_name.")
        for k, v in context.items():
            if not isinstance(v, (list, dict)):
                doc.add_paragraph(f"{k}: {v}", style='List Bullet')

    return doc
